"""Document processing service for extracting and chunking text from uploaded files."""

from __future__ import annotations

import csv
import io
from pathlib import Path
from typing import List, Tuple

import pandas as pd
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

from app.core.config import Settings


class DocumentProcessor:
    """Service for extracting text from various file formats and chunking content."""

    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
    CHUNK_SIZE = 1000  # Characters per chunk
    CHUNK_OVERLAP = 200  # Overlap between chunks

    SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".csv"}

    def __init__(self, settings: Settings) -> None:
        """Initialize the document processor."""

        self._settings = settings

    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """Validate file size and extension."""

        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            return False, f"File size exceeds {self.MAX_FILE_SIZE / (1024*1024):.1f} MB limit"

        # Check extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"

        return True, ""

    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text content from uploaded file based on file type."""

        file_ext = Path(filename).suffix.lower()

        if file_ext == ".pdf":
            return self._extract_from_pdf(file_content)
        elif file_ext in {".doc", ".docx"}:
            return self._extract_from_docx(file_content)
        elif file_ext == ".csv":
            return self._extract_from_csv(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")

    def chunk_text(self, text: str, chunk_size: int | None = None, overlap: int | None = None) -> List[str]:
        """Split text into overlapping chunks for embedding."""

        chunk_size = chunk_size or self.CHUNK_SIZE
        overlap = overlap or self.CHUNK_OVERLAP

        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings in the last 100 chars
                last_period = chunk.rfind(". ")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)

                if break_point > chunk_size - 200:  # Don't break too early
                    chunk = chunk[: break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())
            start = end - overlap  # Overlap for context continuity

        return chunks

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""

        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"Failed to extract text from PDF: {str(exc)}") from exc

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""

        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"Failed to extract text from DOCX: {str(exc)}") from exc

    def _extract_from_csv(self, file_content: bytes) -> str:
        """Extract text from CSV file."""

        try:
            csv_file = io.BytesIO(file_content)
            # Try to detect encoding
            try:
                df = pd.read_csv(csv_file, encoding="utf-8")
            except UnicodeDecodeError:
                csv_file.seek(0)
                df = pd.read_csv(csv_file, encoding="latin-1")

            # Convert DataFrame to readable text format
            text_parts = []

            # Add headers
            headers = " | ".join(str(col) for col in df.columns)
            text_parts.append(f"Headers: {headers}")

            # Add rows (limit to prevent huge outputs)
            max_rows = 1000
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join(str(val) for val in row.values if pd.notna(val))
                text_parts.append(f"Row {idx + 1}: {row_text}")

            if len(df) > max_rows:
                text_parts.append(f"\n... ({len(df) - max_rows} more rows)")

            return "\n".join(text_parts)
        except Exception as exc:
            raise ValueError(f"Failed to extract text from CSV: {str(exc)}") from exc

